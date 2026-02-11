"""
DOCX解析服务
- 将DOCX转换为PDF
- 提取文档结构
"""

import subprocess
import os
from pathlib import Path
from typing import List, Optional
from docx import Document

from models.schemas import PageInfo
from services.pdf_parser import PDFParser


class DocxParser:
    """DOCX文档解析器"""

    def __init__(self):
        self.pdf_parser = PDFParser()

    def parse_to_pdf(self, docx_path: str, output_pdf_path: str) -> List[PageInfo]:
        """
        将DOCX转换为PDF，然后解析
        优先使用LibreOffice，备选其他方法
        """
        docx_path = Path(docx_path)
        output_pdf_path = Path(output_pdf_path)

        # 确保输出目录存在
        output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # 尝试使用LibreOffice转换
        success = self._convert_with_libreoffice(str(docx_path), str(output_pdf_path))

        if not success:
            # 备选：使用python-docx + reportlab生成基础PDF
            success = self._convert_with_python(docx_path, output_pdf_path)

        if not success:
            raise Exception("DOCX转PDF失败")

        # 使用PDF解析器解析转换后的PDF
        return self.pdf_parser.parse(str(output_pdf_path))

    def _convert_with_libreoffice(self, docx_path: str, output_pdf_path: str) -> bool:
        """
        使用LibreOffice命令行转换DOCX到PDF
        """
        try:
            output_dir = Path(output_pdf_path).parent

            # LibreOffice命令
            cmd = [
                'soffice',
                '--headless',
                '--convert-to', 'pdf',
                '--outdir', str(output_dir),
                docx_path
            ]

            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60
            )

            if result.returncode == 0:
                # LibreOffice生成的PDF文件名与原文档相同，只是扩展名不同
                expected_pdf = output_dir / (Path(docx_path).stem + '.pdf')
                if expected_pdf.exists():
                    # 重命名为目标路径
                    if str(expected_pdf) != output_pdf_path:
                        expected_pdf.rename(output_pdf_path)
                    return True

            return False

        except FileNotFoundError:
            print("LibreOffice未安装，尝试其他方法")
            return False
        except subprocess.TimeoutExpired:
            print("LibreOffice转换超时")
            return False
        except Exception as e:
            print(f"LibreOffice转换失败: {e}")
            return False

    def _convert_with_python(self, docx_path: Path, output_pdf_path: Path) -> bool:
        """
        使用Python库转换DOCX到PDF（基础实现）
        注意：此方法生成的PDF格式较简单，仅用于备用
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 注册中文字体（尝试常见字体）
            font_paths = [
                '/System/Library/Fonts/PingFang.ttc',  # macOS
                '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # Linux
                'C:/Windows/Fonts/simhei.ttf',  # Windows
            ]

            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path))
                        font_registered = True
                        break
                    except:
                        continue

            # 创建PDF
            c = canvas.Canvas(str(output_pdf_path), pagesize=A4)
            width, height = A4

            if font_registered:
                c.setFont('Chinese', 12)
            else:
                c.setFont('Helvetica', 12)

            # 读取DOCX内容
            doc = Document(str(docx_path))

            y = height - 50
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    c.drawString(50, y, text[:100])  # 限制长度
                    y -= 20

                    if y < 50:
                        c.showPage()
                        if font_registered:
                            c.setFont('Chinese', 12)
                        y = height - 50

            c.save()
            return True

        except Exception as e:
            print(f"Python转换失败: {e}")
            return False

    def extract_text(self, docx_path: str) -> str:
        """
        提取DOCX文档的所有文本
        """
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return '\n'.join(paragraphs)

    def extract_tables(self, docx_path: str) -> List[List[List[str]]]:
        """
        提取DOCX中的所有表格
        返回: [[[cell1, cell2], [cell3, cell4]], ...]
        """
        doc = Document(docx_path)
        tables = []

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        return tables
